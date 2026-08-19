"""Template-filling toolkit for the SEP490 reports.

Every report is rebuilt from the pristine FPT template rather than patched in place,
because repeated patching is what accumulated the structural damage (body text welded
into headings, orphaned sections, figures numbered out of order) in the first place.

The one rule that matters here, learned the hard way: `clear_regions` for the WHOLE
document runs first, and only then does anything get inserted. `remove_between(A, B)`
deletes whatever currently sits between two headings — it cannot tell "leftover template
instruction" from "content I inserted a moment ago" — so interleaving clear and insert
silently eats the new content.
"""
from __future__ import annotations

import copy
import os

import docx
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image

TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                            "..", "extracted", "SEP490_Report-Document-Template")
OUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "generated")

# A4 with the template's margins: keep every picture inside this box or Word rescales it.
MAX_PIC_W_IN = 6.20
MAX_PIC_H_IN = 8.10

# ── house style ──────────────────────────────────────────────────────────────
# Applied uniformly by the helpers below so presentation never varies section to
# section: justified body text, one consistent paragraph gap, bullets and labelled
# items on a hanging indent so wrapped lines align under the text rather than under
# the label, and headings that cannot be orphaned at the foot of a page.
BODY_SPACE_AFTER_PT = 6
BODY_LINE_SPACING = 1.15
BULLET_INDENT_IN = 0.25
BULLET_HANG_IN = 0.25
LABEL_INDENT_IN = 0.80      # width reserved for the "FE-01:" style label column
TABLE_FONT_PT = 10
CELL_PAD_DXA = 72           # ~0.05 in of breathing room inside every cell


def iter_blocks(doc):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


class Report:
    def __init__(self, template_name: str, out_name: str):
        self.doc = docx.Document(os.path.join(TEMPLATE_DIR, template_name))
        self.out_path = os.path.abspath(os.path.join(OUT_DIR, out_name))
        self._fig_no = 0
        self._fig_index: list[tuple[int, str]] = []

    # ── locating ─────────────────────────────────────────────────────────────
    def para(self, prefix: str, style: str | None = None) -> Paragraph:
        """First paragraph whose text starts with `prefix` (optionally of a given style)."""
        for p in self.doc.paragraphs:
            if p.text.strip().startswith(prefix) and (style is None or p.style.name == style):
                return p
        raise LookupError(f"paragraph not found: {prefix!r}")

    def heading(self, prefix: str) -> Paragraph:
        for p in self.doc.paragraphs:
            if p.style.name.startswith("Heading") and p.text.strip().startswith(prefix):
                return p
        raise LookupError(f"heading not found: {prefix!r}")

    # ── clearing (PASS 1 — must complete before any insert) ──────────────────
    def clear_between(self, start_prefix: str, end_prefix: str) -> int:
        """Delete every block strictly between two headings, keeping both headings."""
        start = self.heading(start_prefix)._p
        end = self.heading(end_prefix)._p
        removed, cur = 0, start.getnext()
        while cur is not None and cur is not end:
            nxt = cur.getnext()
            if cur.tag == qn("w:sectPr"):
                break
            cur.getparent().remove(cur)
            removed += 1
            cur = nxt
        return removed

    def clear_after(self, start_prefix: str) -> int:
        """Delete everything after a heading to the end of the body."""
        start = self.heading(start_prefix)._p
        removed, cur = 0, start.getnext()
        while cur is not None:
            nxt = cur.getnext()
            if cur.tag == qn("w:sectPr"):
                break
            cur.getparent().remove(cur)
            removed += 1
            cur = nxt
        return removed

    def clear_regions(self, *pairs) -> int:
        """PASS 1. Every pair is (start_prefix, end_prefix); end None means to-the-end."""
        total = 0
        for start, end in pairs:
            total += self.clear_after(start) if end is None else self.clear_between(start, end)
        return total

    # ── inserting (PASS 2) ───────────────────────────────────────────────────
    def _style_source(self, style: str):
        """A paragraph of the wanted style to clone formatting from, if the template has one."""
        for p in self.doc.paragraphs:
            if p.style.name == style and p.runs:
                return p
        return None

    @staticmethod
    def _anchor_element(anchor):
        """Accept a Paragraph or a Table as an insertion anchor.

        A table is not a paragraph, so anchoring the next block to
        `doc.paragraphs[-1]` after inserting one silently places it *before* that
        table instead of after it — which is how section content ended up out of
        document order.
        """
        return anchor._tbl if isinstance(anchor, Table) else anchor._p

    def add_after(self, anchor, text: str, style: str = "Normal") -> Paragraph:
        anchor_el = self._anchor_element(anchor)
        src = self._style_source(style)
        if src is not None:
            el = copy.deepcopy(src._p)
            anchor_el.addnext(el)
            par = Paragraph(el, src._parent)
            for r in par.runs[1:]:
                r._element.getparent().remove(r._element)
            par.runs[0].text = text
            par.runs[0].bold = None
            par.runs[0].italic = None
        else:
            el = copy.deepcopy(anchor_el)
            anchor_el.addnext(el)
            par = Paragraph(el, anchor._parent)
            for r in list(par.runs):
                r._element.getparent().remove(r._element)
            par.add_run(text)
            par.style = self.doc.styles[style]
        par.style = self.doc.styles[style]
        self._style_body(par, style)
        return par

    @staticmethod
    def _style_body(par, style: str):
        """One consistent look for every inserted paragraph.

        Body copy is justified with a single fixed gap after it; list items sit on a
        hanging indent so a wrapped line lines up under the text and not under the
        bullet. Doing this here rather than per call site is what keeps the whole
        document even.
        """
        pf = par.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(BODY_SPACE_AFTER_PT)
        pf.line_spacing = BODY_LINE_SPACING
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.widow_control = True
        if style == "List Paragraph":
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.left_indent = Inches(BULLET_INDENT_IN + BULLET_HANG_IN)
            pf.first_line_indent = Inches(-BULLET_HANG_IN)
        else:
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.left_indent = Inches(0)
            pf.first_line_indent = Inches(0)

    def add_labelled(self, anchor: Paragraph, label: str, text: str) -> Paragraph:
        """An "FE-01:  <text>" item whose wrapped lines align under the text, not the label."""
        par = self.add_after(anchor, f"{label}	{text}", "List Paragraph")
        pf = par.paragraph_format
        pf.left_indent = Inches(LABEL_INDENT_IN)
        pf.first_line_indent = Inches(-LABEL_INDENT_IN)
        pf.tab_stops.add_tab_stop(Inches(LABEL_INDENT_IN))
        return par

    def add_paragraphs(self, anchor: Paragraph, texts, style: str = "Normal") -> Paragraph:
        for t in texts:
            anchor = self.add_after(anchor, t, style)
        return anchor

    def add_bullets(self, anchor: Paragraph, items) -> Paragraph:
        for t in items:
            anchor = self.add_after(anchor, "•  " + t, "List Paragraph")
        return anchor

    # ── tables ───────────────────────────────────────────────────────────────
    def add_table(self, anchor, header, rows, widths=None) -> Table:
        """Clone a template table so borders/fonts match, then resize and fill it."""
        src = self._table_template(len(header))
        el = copy.deepcopy(src._tbl)
        self._anchor_element(anchor).addnext(el)
        tbl = Table(el, src._parent)

        self._set_column_count(tbl, len(header))
        # Keep a body row as the prototype for data rows. Cloning the header instead
        # would carry its shading and bold onto every row in the table.
        body_proto = copy.deepcopy(tbl.rows[1]._tr) if len(tbl.rows) > 1 else None
        while len(tbl.rows) > 1:
            tbl._tbl.remove(tbl.rows[-1]._tr)
        self._fill_row(tbl.rows[0], header, bold=True)
        for values in rows:
            tr = copy.deepcopy(body_proto) if body_proto is not None                 else copy.deepcopy(tbl.rows[0]._tr)
            tbl._tbl.append(tr)
            self._fill_row(tbl.rows[-1], values, bold=False)
            self._plain_body_row(tbl.rows[-1])
        if widths:
            self._set_widths(tbl, widths)
        self._style_table(tbl)
        return tbl

    @staticmethod
    def _plain_body_row(row):
        """Force a data row to plain black-on-white.

        The header keeps the template's tinted fill; body rows must not, whatever the
        row they were cloned from happened to carry.
        """
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            for shd in tcPr.findall(qn("w:shd")):
                tcPr.remove(shd)
            tcPr.append(tcPr.makeelement(qn("w:shd"), {
                qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "FFFFFF"}))
            for par in cell.paragraphs:
                for run in par.runs:
                    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                    run.bold = False

    @staticmethod
    def _style_table(tbl):
        """Fixed layout, repeating header, even padding, text top-aligned in every cell.

        Without a fixed layout Word re-flows column widths to suit the longest cell,
        which is what made previous tables look ragged from one page to the next.
        """
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        tblPr = tbl._tbl.tblPr

        layout = tblPr.makeelement(qn("w:tblLayout"), {qn("w:type"): "fixed"})
        tblPr.append(layout)

        margins = tblPr.makeelement(qn("w:tblCellMar"), {})
        for side in ("top", "left", "bottom", "right"):
            node = margins.makeelement(qn(f"w:{side}"),
                                       {qn("w:w"): str(CELL_PAD_DXA), qn("w:type"): "dxa"})
            margins.append(node)
        tblPr.append(margins)

        # Repeat the header row at the top of every page the table spills onto.
        trPr = tbl.rows[0]._tr.get_or_add_trPr()
        trPr.append(trPr.makeelement(qn("w:tblHeader"), {}))

        for i, row in enumerate(tbl.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                for par in cell.paragraphs:
                    pf = par.paragraph_format
                    pf.space_before = Pt(1)
                    pf.space_after = Pt(1)
                    pf.line_spacing = 1.0
                    pf.left_indent = Inches(0)
                    pf.first_line_indent = Inches(0)
                    pf.alignment = (WD_ALIGN_PARAGRAPH.CENTER if i == 0
                                    else WD_ALIGN_PARAGRAPH.LEFT)
                    for run in par.runs:
                        run.font.size = Pt(TABLE_FONT_PT)

    @staticmethod
    def _set_column_count(tbl, want: int):
        """Grow or shrink a cloned template table to exactly `want` columns.

        The template only ships tables of a few widths, so a cloned one rarely has the
        column count a given section needs. Without this, extra columns silently kept the
        template's own header text and requested columns were dropped off the right-hand
        edge.
        """
        grid = tbl._tbl.find(qn("w:tblGrid"))
        cols = grid.findall(qn("w:gridCol")) if grid is not None else []
        have = len(cols)
        if have == want or grid is None:
            return
        if have > want:
            for col in cols[want:]:
                grid.remove(col)
            for tr in tbl._tbl.findall(qn("w:tr")):
                cells = tr.findall(qn("w:tc"))
                for tc in cells[want:]:
                    tr.remove(tc)
        else:
            for _ in range(want - have):
                grid.append(copy.deepcopy(cols[-1]))
                for tr in tbl._tbl.findall(qn("w:tr")):
                    cells = tr.findall(qn("w:tc"))
                    tr.append(copy.deepcopy(cells[-1]))

    def _table_template(self, ncols: int):
        for t in self.doc.tables:
            if len(t.columns) == ncols:
                return t
        return self.doc.tables[0]

    @staticmethod
    def _fill_row(row, values, bold: bool):
        for cell, value in zip(row.cells, values):
            for extra in cell.paragraphs[1:]:
                extra._p.getparent().remove(extra._p)
            par = cell.paragraphs[0]
            for r in par.runs[1:]:
                r._element.getparent().remove(r._element)
            if par.runs:
                par.runs[0].text = str(value)
                par.runs[0].bold = bold
            else:
                par.add_run(str(value)).bold = bold

    @staticmethod
    def _set_widths(tbl, widths_in):
        tbl.autofit = False
        grid = tbl._tbl.find(qn("w:tblGrid"))
        if grid is not None:
            for col, w in zip(grid.findall(qn("w:gridCol")), widths_in):
                col.set(qn("w:w"), str(int(Inches(w).twips)))
        for row in tbl.rows:
            for cell, w in zip(row.cells, widths_in):
                cell.width = Inches(w)

    # ── figures ──────────────────────────────────────────────────────────────
    def add_figure(self, anchor: Paragraph, image_path: str, caption: str) -> Paragraph:
        """Insert a picture scaled to fit the text block, then its numbered caption.

        Figure numbers are issued in document order by this method alone, so a figure
        inserted or removed in the middle can never leave a gap or a duplicate.
        """
        with Image.open(image_path) as im:
            w_px, h_px = im.size
        w_in = MAX_PIC_W_IN
        h_in = w_in * h_px / w_px
        if h_in > MAX_PIC_H_IN:
            h_in = MAX_PIC_H_IN
            w_in = h_in * w_px / h_px

        pic_par = self.add_after(anchor, "", "Normal")
        for r in list(pic_par.runs):
            r._element.getparent().remove(r._element)
        pic_par.add_run().add_picture(image_path, width=Inches(w_in), height=Inches(h_in))
        pf = pic_par.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.left_indent = Inches(0)
        pf.first_line_indent = Inches(0)
        pf.space_before = Pt(6)
        pf.space_after = Pt(2)
        pf.keep_with_next = True          # caption never separates from its figure

        self._fig_no += 1
        text = f"Figure {self._fig_no} — {caption}"
        self._fig_index.append((self._fig_no, caption))
        cap_par = self.add_after(pic_par, text, "Normal")
        cpf = cap_par.paragraph_format
        cpf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cpf.left_indent = Inches(0)
        cpf.first_line_indent = Inches(0)
        cpf.space_before = Pt(0)
        cpf.space_after = Pt(10)
        for r in cap_par.runs:
            r.italic = True
            r.font.size = Pt(10)
        return cap_par

    def figure_number(self, caption_fragment: str) -> int:
        for no, cap in self._fig_index:
            if caption_fragment.lower() in cap.lower():
                return no
        raise LookupError(f"no figure captioned like {caption_fragment!r}")

    # ── record of changes ────────────────────────────────────────────────────
    def record_of_changes(self, rows):
        tbl = self.doc.tables[0]
        # Same rule as every other table: clone a body row, not the tinted header.
        body_proto = copy.deepcopy(tbl.rows[1]._tr) if len(tbl.rows) > 1 else None
        while len(tbl.rows) > 1:
            tbl._tbl.remove(tbl.rows[-1]._tr)
        for values in rows:
            tr = copy.deepcopy(body_proto) if body_proto is not None                 else copy.deepcopy(tbl.rows[0]._tr)
            tbl._tbl.append(tr)
            self._fill_row(tbl.rows[-1], values, bold=False)
            self._plain_body_row(tbl.rows[-1])
        self._style_table(tbl)

    # ── saving ───────────────────────────────────────────────────────────────
    def tidy_headings(self):
        """Even spacing above and below every heading, and never orphaned at a page foot."""
        for par in self.doc.paragraphs:
            name = par.style.name
            if not name.startswith("Heading"):
                continue
            pf = par.paragraph_format
            pf.keep_with_next = True
            pf.page_break_before = False
            pf.space_before = Pt({"Heading 1": 18, "Heading 2": 14,
                                  "Heading 3": 10, "Heading 4": 8}.get(name, 8))
            pf.space_after = Pt(6)
            pf.left_indent = Inches(0)
            pf.first_line_indent = Inches(0)

    def save(self):
        self.tidy_headings()
        # Force Word to rebuild the table of contents on open; the template ships a
        # stale cached rendering of it.
        settings = self.doc.settings.element
        uf = settings.find(qn("w:updateFields"))
        if uf is None:
            uf = settings.makeelement(qn("w:updateFields"), {})
            settings.append(uf)
        uf.set(qn("w:val"), "true")

        os.makedirs(os.path.dirname(self.out_path), exist_ok=True)
        self.doc.save(self.out_path)
        return self.out_path
