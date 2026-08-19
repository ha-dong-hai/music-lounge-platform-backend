"""Automated quality gate for the generated reports.

Encodes the document/diagram error taxonomy the checks are meant to catch, so a defect is
found by running this rather than by someone spotting it while reading:

  layout      — an image wider or taller than the text block, so Word rescales and distorts it
  numbering   — figure numbers that duplicate, skip, or disagree with an in-text reference
  captioning  — an image with no caption, or a caption with no image
  structure   — an empty heading, or body prose accidentally carrying a heading style
  tables      — a row whose cell count disagrees with its header; body cells not black on white
  leftovers   — template instruction text, sample content, or unfilled placeholders
  consistency — the same quantity stated differently in two reports
  render      — a package with duplicate parts, or a stale table of contents

Exit code is non-zero if any FAIL is reported, so it can gate a build.
"""
from __future__ import annotations

import os
import re
import sys
import zipfile

import docx
from docx.oxml.ns import qn
from docx.shared import Emu
from docx.table import Table
from docx.text.paragraph import Paragraph

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import facts as F

GEN = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "generated"))
MAX_W_IN, MAX_H_IN = 6.25, 8.15
CAPTION = re.compile(r"^Figure (\d+)\s+—")
INSTRUCTION = re.compile(r"\[[A-Z][^\]]{30,}\]")
SAMPLE_WORDS = ["Cafeteria", "Process Impact", "SWP493", "KhanhNTHE", "Member A",
                "<Feature Name", "<Function Name", "System name1", "Workflow 1</"]

failures: list[str] = []
warnings: list[str] = []


def fail(doc, msg):
    failures.append(f"{doc}: {msg}")


def warn(doc, msg):
    warnings.append(f"{doc}: {msg}")


def blocks(d):
    for ch in d.element.body.iterchildren():
        if ch.tag == qn("w:p"):
            yield Paragraph(ch, d)
        elif ch.tag == qn("w:tbl"):
            yield Table(ch, d)


def all_text(d):
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                parts.append(c.text)
    return " ".join(parts)


def check_package(name, path):
    names = zipfile.ZipFile(path).namelist()
    dupes = sorted({n for n in names if names.count(n) > 1})
    if dupes:
        fail(name, f"package has duplicate parts: {dupes[:4]}")


def check_toc(name, d):
    uf = d.settings.element.find(qn("w:updateFields"))
    if uf is None or uf.get(qn("w:val")) not in ("true", "1"):
        fail(name, "table of contents will not refresh on open (updateFields not set)")


def check_figures(name, d):
    seq = []
    for blk in blocks(d):
        if not isinstance(blk, Paragraph):
            continue
        has_img = bool(blk._p.findall(".//" + qn("w:drawing")))
        m = CAPTION.match(blk.text.strip())
        if has_img:
            seq.append(("img", blk))
        elif m:
            seq.append(("cap", int(m.group(1))))

    caps = [n for kind, n in seq if kind == "cap"]
    if caps != list(range(1, len(caps) + 1)):
        fail(name, f"figure numbers are not sequential from 1: {caps}")

    # every caption must be immediately preceded by an image
    for i, (kind, val) in enumerate(seq):
        if kind == "cap" and (i == 0 or seq[i - 1][0] != "img"):
            fail(name, f"Figure {val} caption has no image directly above it")
    imgs_with_caption = sum(1 for i, (k, _) in enumerate(seq)
                            if k == "img" and i + 1 < len(seq) and seq[i + 1][0] == "cap")
    uncaptioned = sum(1 for k, _ in seq if k == "img") - imgs_with_caption
    if uncaptioned > 1:  # the cover-page logo is allowed to be uncaptioned
        warn(name, f"{uncaptioned} images have no caption")

    # in-text references must point at a figure that exists
    for p in d.paragraphs:
        text = p.text.strip()
        if CAPTION.match(text):
            continue
        for ref in re.findall(r"Figure (\d+)", text):
            if caps and int(ref) > len(caps):
                fail(name, f"text references Figure {ref} but only {len(caps)} figures exist")


def check_images(name, d):
    for sh in d.inline_shapes:
        w, h = Emu(sh.width).inches, Emu(sh.height).inches
        if w > MAX_W_IN or h > MAX_H_IN:
            fail(name, f"image {w:.2f}x{h:.2f} in exceeds the text block "
                       f"({MAX_W_IN}x{MAX_H_IN}) and will be rescaled")


def check_tables(name, d):
    for ti, t in enumerate(d.tables):
        ncols = len(t.columns)
        for ri, row in enumerate(t.rows):
            if len(row.cells) != ncols:
                fail(name, f"table {ti} row {ri} has {len(row.cells)} cells but {ncols} columns")
            if ri == 0:
                continue
            for c in row.cells:
                tcPr = c._tc.tcPr
                shd = tcPr.find(qn("w:shd")) if tcPr is not None else None
                fill = shd.get(qn("w:fill")) if shd is not None else None
                if fill not in (None, "FFFFFF", "auto"):
                    fail(name, f"table {ti} row {ri} body cell is shaded {fill}, not white")
                    break
        if len(t.rows) < 2:
            warn(name, f"table {ti} has a header but no data rows")


def check_structure(name, d):
    for p in d.paragraphs:
        style = p.style.name
        text = p.text.strip()
        if style.startswith("Heading"):
            if not text:
                fail(name, f"empty {style}")
            elif len(text) > 120:
                fail(name, f"{style} holds body-length text ({len(text)} chars): {text[:60]}…")


def check_leftovers(name, d):
    text = all_text(d)
    for m in INSTRUCTION.findall(text):
        fail(name, f"template instruction text remains: {m[:60]}…")
    for w in SAMPLE_WORDS:
        if w in text:
            fail(name, f"template sample content remains: {w!r}")
    ph = sorted(set(re.findall(r"<<[^<>]{1,60}>>", text)))
    if ph:
        warn(name, f"unfilled placeholders (expected — team contact details): {ph}")

    # Single-angle and bracketed markers are the template's own placeholder styles.
    # An earlier revision only looked for the << >> form and so missed an entirely
    # unfilled cover page.
    stripped = re.sub(r"<<[^<>]{1,60}>>", "", text)
    single = sorted({m for m in re.findall(r"<[A-Za-z][^<>]{2,40}>", stripped)
                     if m not in ("<Feature>",)})   # Application.<Feature> is real notation
    for m in single:
        fail(name, f"unfilled template placeholder: {m}")
    for m in sorted(set(re.findall(r"\[[A-Z][a-z][^\]]{2,30}\]", text))):
        fail(name, f"unfilled template placeholder: {m}")


def check_consistency(docs):
    """The same quantity must read the same in every report that states it."""
    S, T = F.SCALE, F.TESTS
    expectations = {
        "effort": (str(S["effort_man_days"]), r"(\d+) man-days"),
        "use cases": (str(S["use_cases"]), r"(\d+) use cases"),
        "entities": (str(S["entities"]), r"(\d+) entities"),
        "total tests": (str(T["total_tests"]), r"(\d+) tests"),
    }
    for label, (expected, pattern) in expectations.items():
        seen = {}
        for name, d in docs.items():
            for found in set(re.findall(pattern, all_text(d))):
                seen.setdefault(found, []).append(name)
        wrong = {v: docs_ for v, docs_ in seen.items() if v != expected}
        if wrong:
            fail("cross-report", f"{label}: expected {expected} but found {wrong}")


def main():
    files = sorted(f for f in os.listdir(GEN) if f.endswith(".docx") and not f.startswith("~"))
    if not files:
        print("no generated reports found")
        return 1

    docs = {}
    print(f"Verifying {len(files)} report(s)\n")
    for f in files:
        path = os.path.join(GEN, f)
        short = f.split("_")[0]
        d = docx.Document(path)
        docs[short] = d

        check_package(short, path)
        check_toc(short, d)
        check_figures(short, d)
        check_images(short, d)
        check_tables(short, d)
        check_structure(short, d)
        check_leftovers(short, d)

        caps = sum(1 for p in d.paragraphs if CAPTION.match(p.text.strip()))
        print(f"  {short:9s} {len(d.paragraphs):4d} paras  {len(d.tables):3d} tables  "
              f"{len(d.inline_shapes):3d} images  {caps:2d} figures")

    check_consistency(docs)

    print()
    if warnings:
        print(f"WARNINGS ({len(warnings)})")
        for w in warnings:
            print(f"  ! {w}")
        print()
    if failures:
        print(f"FAILURES ({len(failures)})")
        for f_ in failures:
            print(f"  x {f_}")
        return 1
    print("PASS — no failures")
    return 0


if __name__ == "__main__":
    sys.exit(main())
