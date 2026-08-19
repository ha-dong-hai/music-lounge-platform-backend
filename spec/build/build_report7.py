"""Report 7 — Final Project Report.

A consolidation, so its content is copied from the freshly built Reports 1-6 rather than
written again. Writing it separately is what previously let it drift from its sources.

Two things must be handled on the way in, both of which bit us before:
  * images — re-added from their bytes so the destination package assigns a fresh part
    name; relating straight to the source part reuses word/media/imageN.png and produces
    a package with duplicate entries;
  * figure numbers — each source report numbers its own figures from 1, so after merging
    they are renumbered once, in document order, together with the in-text references.
"""
import copy
import re
import sys, os
from io import BytesIO

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import docx
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from docxkit import Report, OUT_DIR
import facts as F

S, T = F.SCALE, F.TESTS

PARTS = [
    ("I. Project Introduction", "Report1_Project Introduction - MusicLounge.docx"),
    ("II. Project Management Plan", "Report2_Project Management Plan - MusicLounge.docx"),
    ("III. Software Requirement Specification", "Report3_Software Requirement Specification - MusicLounge.docx"),
    ("IV. Software Design Description", "Report4_Software Design Document - MusicLounge.docx"),
    ("V. Software Testing Documentation", "Report5_Test Documentation - MusicLounge.docx"),
    ("VI. Release Package & User Guides", "Report6_Software User Guides - MusicLounge.docx"),
]

r = Report("Report7_Final Project Report.docx",
           "Report7_Final Project Report - MusicLounge.docx")
dst = r.doc


def part_headings():
    return [(p.text.strip(), p._p) for p in dst.paragraphs
            if p.style.name == "Heading 1" and p.text.strip()]


def source_content(src_doc):
    """Every block after the source report's 'II. ...' heading — its real content."""
    children = list(src_doc.element.body.iterchildren())
    start = None
    for i, ch in enumerate(children):
        if ch.tag == qn("w:p"):
            par = Paragraph(ch, src_doc)
            if par.style.name == "Heading 1" and par.text.strip().startswith("II."):
                start = i + 1
                break
    assert start is not None, "source has no 'II.' heading"
    return [c for c in children[start:] if c.tag != qn("w:sectPr")]


def remap_images(el, src_doc):
    for blip in el.findall(".//" + qn("a:blip")):
        rid = blip.get(qn("r:embed"))
        if not rid:
            continue
        blob = src_doc.part.rels[rid].target_part.blob
        new_rid, _ = dst.part.get_or_add_image(BytesIO(blob))
        blip.set(qn("r:embed"), new_rid)


# ── PASS 1: clear every Part ─────────────────────────────────────────────────
for title, _ in PARTS:
    headings = part_headings()
    names = [n for n, _ in headings]
    i = names.index(title)
    start = headings[i][1]
    end = headings[i + 1][1] if i + 1 < len(headings) else None
    cur = start.getnext()
    while cur is not None and cur is not end:
        nxt = cur.getnext()
        if cur.tag == qn("w:sectPr"):
            break
        cur.getparent().remove(cur)
        cur = nxt

# also clear the two front-matter sections
for start_t, end_t in [("Acknowledgement", "Definition and Acronyms"),
                       ("Definition and Acronyms", "I. Project Introduction")]:
    headings = dict(part_headings())
    cur = headings[start_t].getnext()
    stop = headings[end_t]
    while cur is not None and cur is not stop:
        nxt = cur.getnext()
        cur.getparent().remove(cur)
        cur = nxt

# ── PASS 2: front matter ─────────────────────────────────────────────────────
ack = r.heading("Acknowledgement")
r.add_paragraphs(ack, [
    "We would like to thank our supervisor, "
    f"{F.SUPERVISOR['name']}, for the guidance and the direct, practical feedback that shaped this "
    "project at every milestone — particularly for pressing us on the difference between a system "
    "that demonstrates well and one that would survive real use.",
    "We also thank the lecturers of the Software Engineering programme at FPT University, whose "
    "courses in software architecture, database design, testing and project management gave us the "
    "foundations this work builds on.",
    "Finally, we thank the small live-music venues whose day-to-day practice this platform is "
    "modelled on. Understanding how they really sell tickets, run a night and pay performers is what "
    "kept the design honest.",
])

acr = r.heading("Definition and Acronyms")
r.add_table(acr, ["Acronym / Term", "Definition"], [
    ["API", "Application Programming Interface — the contract every client application calls."],
    ["ASVS", "Application Security Verification Standard, published by OWASP; used here as the security checklist."],
    ["CDN", "Content Delivery Network — edge caching used to serve images without loading the API."],
    ["CI/CD", "Continuous Integration and Continuous Deployment — the automated build, test and release pipeline."],
    ["CQRS", "Command Query Responsibility Segregation — the pattern separating write commands from read queries."],
    ["Crow's-foot", "The entity-relationship notation used in this report to express cardinality."],
    ["DTO", "Data Transfer Object — the shape returned to a client, distinct from the stored entity."],
    ["End-to-end (E2E)", "A test that drives a complete journey across frontend, API and database together."],
    ["Entity", "A business object persisted as a database table."],
    ["Golden path", "The main successful route an actor takes through a feature, used as the acceptance scenario."],
    ["Hold", "A temporary reservation of ticket capacity taken before payment and released if unpaid."],
    ["IPN", "Instant Payment Notification — the payment gateway's server-to-server confirmation callback."],
    ["Ledger (double-entry)", "Accounting record in which every transaction posts balancing debit and credit lines."],
    ["Lounge / Venue", "A physical live-music venue registered on the platform."],
    ["OWASP", "Open Worldwide Application Security Project — the source of the security standard applied."],
    ["Performer", "An artist appearing in a show's line-up; a catalogue record, not a login."],
    ["Settlement", "A scheduled payout of a venue's share of ticket revenue, released in two tranches."],
    ["Show", "One concrete performance event at a venue on a given evening."],
    ["SLA", "Service Level Agreement — here, the deadline within which a review or response must happen."],
    ["Surface", "One of the five client applications through which the platform is used."],
    ["Tier", "A priced category of ticket access for a show."],
    ["UAT", "User Acceptance Testing — manual walkthrough of each actor's golden path before release."],
    ["UML", "Unified Modeling Language — the notation used for the diagrams in this report."],
    ["WCAG", "Web Content Accessibility Guidelines — the accessibility standard applied to the web surfaces."],
], widths=[1.5, 4.7])

# ── PASS 3: copy each Part's content from its source report ──────────────────
copied_total = 0
for title, filename in PARTS:
    headings = part_headings()
    names = [n for n, _ in headings]
    anchor = headings[names.index(title)][1]

    src = docx.Document(os.path.join(OUT_DIR, filename))
    n = 0
    for el in source_content(src):
        new_el = copy.deepcopy(el)
        remap_images(new_el, src)
        anchor.addnext(new_el)
        anchor = new_el
        n += 1
    copied_total += n
    print(f"  {title:44s} {n:4d} blocks")

# ── PASS 4: renumber figures across the merged document ──────────────────────
CAPTION = re.compile(r"^Figure \d+ —")
captions = [p for p in dst.paragraphs if CAPTION.match(p.text.strip())]
old_to_new = {}
for new_no, par in enumerate(captions, 1):
    for run in par.runs:
        m = re.match(r"\s*Figure (\d+)", run.text)
        if m:
            old_to_new.setdefault((id(par), int(m.group(1))), new_no)
            run.text = re.sub(r"Figure \d+", f"Figure {new_no}", run.text, count=1)
            break

# In-text references. Each source numbered from 1, so a bare "Figure N" in body text is
# resolved against the captions of the Part it sits in.
part_ids = [h[1] for h in part_headings()]


def part_index_of(par):
    prev = par._p.getprevious()
    idx = -1
    node = par._p
    while node is not None:
        if node in part_ids:
            return part_ids.index(node)
        node = node.getprevious()
    return idx


caption_by_part = {}
for new_no, par in enumerate(captions, 1):
    caption_by_part.setdefault(part_index_of(par), []).append(new_no)

for par in dst.paragraphs:
    text = par.text.strip()
    if CAPTION.match(text) or "Figure" not in text:
        continue
    pi = part_index_of(par)
    seq = caption_by_part.get(pi, [])
    if not seq:
        continue
    for run in par.runs:
        if "Figure" not in run.text:
            continue

        def fix(m):
            old = int(m.group(1))
            return f"Figure {seq[old - 1]}" if 1 <= old <= len(seq) else m.group(0)

        run.text = re.sub(r"Figure (\d+)", fix, run.text)

print(f"\n  figures renumbered: {len(captions)}")

# ── PASS 5: cover page ───────────────────────────────────────────────────────
# Report 7 has no Record of Changes table — tables[0] is the cover block. Writing a
# change-log row into it (as the shared helper would) corrupts the cover, so the cover
# is filled explicitly instead.
def set_cell_text(cell, text):
    for extra in cell.paragraphs[1:]:
        extra._p.getparent().remove(extra._p)
    par = cell.paragraphs[0]
    for run in par.runs[1:]:
        run._element.getparent().remove(run._element)
    if par.runs:
        par.runs[0].text = text
    else:
        par.add_run(text)


members = chr(10).join(
    f"{m['name']} — <<RollNo>> — <<Student code>>" for m in F.TEAM)

title_tbl = dst.tables[1]
for row in title_tbl.rows:
    if row.cells[0].text.strip().startswith("[Project name]"):
        set_cell_text(row.cells[0], F.PROJECT["name"])

info_tbl = dst.tables[2]
for row in info_tbl.rows:
    label = row.cells[0].text.strip()
    if label.startswith("<Group Name>"):
        set_cell_text(row.cells[0], F.PROJECT["group"])
        set_cell_text(row.cells[1], F.PROJECT["group"])
    elif label.startswith("Group Members"):
        set_cell_text(row.cells[1], members)
    elif label.startswith("Ext Supervisor"):
        set_cell_text(row.cells[1], "Not applicable")
    elif label.startswith("Supervisor"):
        set_cell_text(row.cells[1], F.SUPERVISOR["name"])
    elif label.startswith("Capstone Project code"):
        set_cell_text(row.cells[1], F.PROJECT["code"])

for par in dst.paragraphs:
    if "<month>" in par.text or "<year>" in par.text:
        joined = par.text.replace("<month>/<year>", "August 2026")
        for extra in par.runs[1:]:
            extra._element.getparent().remove(extra._element)
        if par.runs:
            par.runs[0].text = joined

print("  cover page filled")

path = r.save()
print(f"\nbuilt {path} — {copied_total} blocks copied")
